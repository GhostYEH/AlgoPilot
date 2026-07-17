from copy import deepcopy
from pathlib import Path
import re

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm


DOCX_DIR = Path(__file__).resolve().parents[1] / "docx"
KEEP = {
    "01_AlgoPilot项目说明书.docx",
    "02_AlgoPilot系统开发说明书.docx",
    "03_AlgoPilot测试说明书.docx",
    "05_AlgoPilot用户操作手册.docx",
}


def set_paragraph_text(paragraph, text):
    old_rpr = None
    if paragraph.runs and paragraph.runs[0]._r.rPr is not None:
        old_rpr = deepcopy(paragraph.runs[0]._r.rPr)
    paragraph.clear()
    run = paragraph.add_run(text)
    if old_rpr is not None:
        run._r.insert(0, old_rpr)


def delete_paragraph(paragraph):
    parent = paragraph._element.getparent()
    parent.remove(paragraph._element)
    paragraph._p = paragraph._element = None


def ensure_heading4(doc):
    try:
        style = doc.styles["Heading 4"]
    except KeyError:
        style = doc.styles.add_style("Heading 4", WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = "黑体"
    style.font.size = Pt(11)
    style.font.bold = True
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    pf = style.paragraph_format
    pf.left_indent = Cm(0)
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(6)
    pf.space_after = Pt(3)
    pf.keep_with_next = True
    return style


def add_chinese_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)

    levels = [
        (0, "chineseCounting", "%1、", 720),
        (1, "decimal", "%2.", 1080),
        (2, "decimal", "（%3）", 1440),
    ]
    for ilvl, fmt, label, left in levels:
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(ilvl))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), label)
        lvl.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl.append(suff)
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(left))
        tabs.append(tab)
        ppr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), "360")
        ppr.append(ind)
        lvl.append(ppr)
        rpr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Times New Roman")
        fonts.set(qn("w:hAnsi"), "Times New Roman")
        fonts.set(qn("w:eastAsia"), "宋体")
        rpr.append(fonts)
        lvl.append(rpr)
        abstract.append(lvl)
    numbering.append(abstract)
    return abstract_id


def new_num_instance(doc, abstract_id):
    numbering = doc.part.numbering_part.element
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return num_id


def bullet_level(paragraph):
    name = paragraph.style.name
    if name.endswith(" 3"):
        return 2
    if name.endswith(" 2"):
        return 1
    num_pr = paragraph._p.pPr.numPr if paragraph._p.pPr is not None else None
    if num_pr is not None and num_pr.ilvl is not None:
        return min(int(num_pr.ilvl.val), 2)
    return 0


def apply_number(paragraph, num_id, level):
    ppr = paragraph._p.get_or_add_pPr()
    old = ppr.find(qn("w:numPr"))
    if old is not None:
        ppr.remove(old)
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num_pr.append(ilvl)
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(num)
    ppr.append(num_pr)
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    ind.set(qn("w:left"), str(720 + level * 360))
    ind.set(qn("w:hanging"), "360")


def remove_row(table, row_index):
    table._tbl.remove(table.rows[row_index]._tr)


def remove_column(table, column_index):
    grid = table._tbl.tblGrid
    if grid is not None and len(grid.gridCol_lst) > column_index:
        grid.remove(grid.gridCol_lst[column_index])
    for row in table.rows:
        if len(row.cells) > column_index:
            row._tr.remove(row.cells[column_index]._tc)


def set_version_table_widths(table):
    if not table.rows:
        return
    header = [cell.text.strip() for cell in table.rows[0].cells]
    if header != ["版本", "日期", "说明"]:
        return
    widths = [Cm(2.3), Cm(3.0), Cm(8.2)]
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    grid = table._tbl.tblGrid
    if grid is not None and len(grid.gridCol_lst) == 3:
        for idx, width in enumerate(widths):
            grid.gridCol_lst[idx].set(qn("w:w"), str(int(width.twips)))
    tbl_w = table._tbl.tblPr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        table._tbl.tblPr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(int(width.twips) for width in widths)))
    tbl_w.set(qn("w:type"), "dxa")
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = width
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width.twips)))
            tc_w.set(qn("w:type"), "dxa")


def clean_doc(path):
    doc = Document(path)
    ensure_heading4(doc)
    abstract_id = add_chinese_numbering(doc)

    current_num = None
    previous_bullet = False
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if text.startswith("对应代码版本：") or text.startswith("对应代码版本:"):
            delete_paragraph(paragraph)
            previous_bullet = False
            continue
        if re.match(r"^#{2,6}\s+", text):
            text = re.sub(r"^#{2,6}\s+", "", text)
            set_paragraph_text(paragraph, text)
            paragraph.style = doc.styles["Heading 4"]

        if "[" in paragraph.text and "](" in paragraph.text:
            cleaned = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", paragraph.text)
            if cleaned != paragraph.text:
                set_paragraph_text(paragraph, cleaned)

        is_bullet = paragraph.style.name.startswith("List Bullet")
        if is_bullet:
            if not previous_bullet:
                current_num = new_num_instance(doc, abstract_id)
            apply_number(paragraph, current_num, bullet_level(paragraph))
        previous_bullet = is_bullet

    for table in doc.tables:
        commit_rows = []
        for idx, row in enumerate(table.rows):
            if len(row.cells) == 2 and any(cell.text.strip().lower() == "对应 commit".lower() for cell in row.cells):
                commit_rows.append(idx)
        for idx in reversed(commit_rows):
            remove_row(table, idx)

        if table.rows:
            header = [cell.text.strip().lower() for cell in table.rows[0].cells]
            for idx in reversed(range(len(header))):
                if header[idx] == "对应 commit".lower():
                    remove_column(table, idx)
        set_version_table_widths(table)

    doc.core_properties.author = "AlgoPilot 项目团队"
    doc.core_properties.last_modified_by = "AlgoPilot 项目团队"
    doc.core_properties.comments = ""
    doc.save(path)


def main():
    paths = sorted(p for p in DOCX_DIR.glob("*.docx") if p.name in KEEP)
    if {p.name for p in paths} != KEEP:
        missing = sorted(KEEP - {p.name for p in paths})
        raise SystemExit(f"Expected four retained documents; missing: {missing}")
    for path in paths:
        clean_doc(path)
        print(path)


if __name__ == "__main__":
    main()
