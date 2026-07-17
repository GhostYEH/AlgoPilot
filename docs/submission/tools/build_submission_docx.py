from __future__ import annotations

import re
import zipfile
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[3]
SUB = ROOT / "docs" / "submission"
IMG = SUB / "images"
OUT = SUB / "docx"
OUT.mkdir(parents=True, exist_ok=True)
SHA = "318f21c2bbaf9826dc94e8e31401b247b2c8f6d1"
DATE = "2026-07-17"

DOCS = [
    ("01_项目说明书.md", "01_AlgoPilot项目说明书.docx", "AlgoPilot 项目说明书"),
    ("02_系统开发说明书.md", "02_AlgoPilot系统开发说明书.docx", "AlgoPilot 系统开发说明书"),
    ("03_测试说明书.md", "03_AlgoPilot测试说明书.docx", "AlgoPilot 测试说明书"),
    ("05_用户操作手册.md", "05_AlgoPilot用户操作手册.docx", "AlgoPilot 用户操作手册"),
]

def replace_metadata(text: str) -> str:
    text = re.sub(r"更新日期：\d{4}-\d{2}-\d{2}", f"更新日期：{DATE}", text)
    text = re.sub(r"> 对应代码版本：.*", f"> 对应代码版本：`{SHA}`", text)
    text = text.replace("文档状态：提交候选版", "文档状态：正式提交版")
    text = text.replace("提交前需随最终代码重新确认", SHA)
    return text

def insert_after(text: str, needle: str, block: str) -> str:
    pos = text.find(needle)
    if pos < 0 or block.strip() in text:
        return text
    end = text.find("\n", pos)
    return text[:end+1] + "\n" + block.strip() + "\n" + text[end+1:]

def fig(path: str, caption: str) -> str:
    return f"![{caption}](images/{path})\n\n*{caption}*"

def update_sources():
    for md, _, _ in DOCS:
        p = SUB / md
        text = replace_metadata(p.read_text(encoding="utf-8"))
        text = text.replace("66 个", "67 个").replace("66 端点", "67 端点")
        if md == "03_测试说明书.md":
            text = re.sub(r"> 测试执行环境：.*", "> 测试执行环境：Windows 11 · Python 3.13.7 · Node.js 25.8.1 · npm 11.11.0 · SQLite 测试数据库", text)
            text = re.sub(r"> 真实测试记录（.*", "> 真实测试记录（2026-07-17 执行）：190 passed / 0 failed / 0 skipped / 1 warning / 124.95 秒", text)
            text = text.replace("104.14", "124.95").replace("2026-07-16", DATE)
            text = re.sub(r"\| 测试函数总数 \|.*\|", "| pytest 实际收集并执行 | 190 个 |", text)
            text = text.replace("AlgoPilot 后端测试套件共 **191 个测试函数**（pytest 实际收集 190 个通过）", "AlgoPilot 后端 pytest 本轮实际收集并执行 **190 个测试**")
            text = text.replace("（pytest 实际收集 190 个通过，1 个为条件跳过/未收集）", "")
            text = text.replace("Python 3.13 + Node.js 22.x LTS", "Python 3.13.7 + Node.js 25.8.1 + npm 11.11.0")
            text = text.replace("预期输出：`190 passed, 1 warning in 124.95s`", "本轮实际输出：`190 passed, 1 warning in 124.95s`")
            evidence = """## 本轮冻结验证证据

```text
$ ruff check .
All checks passed!

$ pytest --tb=no -q
190 passed, 1 warning in 124.95s (0:02:04)
```

```text
$ npm run typecheck
vue-tsc -b --noEmit
结果：通过

$ npm run build
3683 modules transformed；built in 3.15s
结果：通过（存在动态导入与大 chunk 提示，不影响构建成功）

$ npm run test:oj-struggle
ojStruggleSession.test.ts: all passed
$ npm run test:path-replan-diff
pathReplanDiff.test.ts: all passed
$ npm run test:graph-module
graphModuleContract.test.ts: all passed
```

> 本轮未执行 pytest-cov，因此不报告覆盖率百分比。"""
            text = insert_after(text, "---", evidence)
        elif md == "04_部署说明书.md":
            text = text.replace(
                "如需单文件分发，可在 `.spec` 中调整为 `onefile=True` 并重新打包，但启动速度会变慢（需先解压到临时目录）。",
                "当前仅保留并验证 COLLECT 目录分发模式。单文件模式未在本轮验证，仅可作为后续研究事项，不提供未经验证的配置步骤。",
            )
            evidence = f"""## 本轮实际启动验证

后端以 `uvicorn main:app --host 127.0.0.1 --port 9000` 启动，前端以 `npm run dev -- --host 127.0.0.1 --port 5173` 启动。`/api/health` 返回 `status=ok`，前端首页返回 HTTP 200。

{fig('DEP-health.png','图 4-1 /api/health 真实返回结果')}

{fig('DEP-fastapi-docs.png','图 4-2 FastAPI Swagger 文档页面')}

{fig('home-system.png','图 4-3 前端成功启动后的系统首页')}

> 本轮仓库中未发现 `backend/dist/AlgoPilot/` 或 `AlgoPilot.exe` 产物，因此未生成打包目录与 EXE 启动截图，也不将其描述为本轮已验证结果。"""
            text = insert_after(text, "---", evidence)
        elif md == "05_用户操作手册.md":
            text = re.sub(r"\n> \*\*截图说明\*\*：.*?\n", "\n", text)
            text = re.sub(r"\n> \*\*待补截图 S\d+：.*?(?=\n\n(?!>)|\n###|\n##)", "", text, flags=re.S)
            text = re.sub(r"<!--.*?-->", "", text, flags=re.S)
            text = re.sub(r"\n## 十六、截图进度表.*\Z", "\n", text, flags=re.S)
            text = text.replace("查看班级学情概览", "查看当前系统实例内学生学习记录的只读聚合视图")
            text = text.replace("显示活跃度统计、OJ 通过率、薄弱模块分布与学习进度分布。", "显示平均掌握度、画像覆盖、OJ 学习事件数、薄弱模块与学习进度分布；不展示无真实时序依据的趋势图或伪造通过率。")
            text = text.replace("### 11.2 班级学情概览", "### 11.2 系统实例学情概览")
            text = text.replace("- **活跃度统计**：学生登录与学习频次\n- **OJ 通过率**：全班各题通过率", "- **画像覆盖**：当前实例内已有画像的学生数\n- **OJ 学习记录数**：成功与失败学习事件合计")
            text = text.replace("### 8.5 学习事件日志\n\n可审计的学习事件记录，包括：", "### 8.5 学习事件审计说明\n\n当前学生端“我的学习”页面未提供独立事件日志列表，不应按原路径截图。学习事件与 Agent 处理记录由后端保存，并可在教师端资源工作台或接口层按权限审计，包括：")
            inserts = [
                ("### 1.1 新用户注册", fig("S01-login-register.png", "图 5-1 注册页面")),
                ("### 2.3 查看画像结果", fig("S03-persona-result.png", "图 5-3 学生画像概览") + "\n\n" + fig("S04-persona-evidence.png", "图 5-4 画像证据链与学习记忆")),
                ("### 3.1 进入学习路径", fig("S05-learning-path.png", "图 5-5 个性化学习路径界面")),
                ("### 3.4 路径重规划", fig("S16-path-replan.png", "图 5-6 真实受挫触发后的路径巩固节点")),
                ("### 4.3 查看生成过程", fig("S06-agent-progress.png", "图 5-7 多智能体资源 DAG 工作台")),
                ("### 4.4 使用生成的资源", fig("S07-generated-resources.png", "图 5-8 已生成资源库")),
                ("### 6.1 进入 OJ", fig("S08-oj-list.png", "图 5-9 OJ 题目列表")),
                ("### 6.2 编写代码", fig("S09-oj-workbench.png", "图 5-10 OJ 题目与代码工作台")),
                ("### 7.1 进入 Trace", fig("S10-trace-visualization.png", "图 5-11 真实执行轨迹概览")),
                ("### 7.5 AI 深度诊断", fig("S11-ai-diagnosis.png", "图 5-12 真实 WA 复现与规则诊断报告")),
                ("### 8.1 我的学习", fig("S12-learning-dashboard.png", "图 5-13 我的学习数据中心")),
                ("### 8.4 学习记忆", fig("S13-learning-memory.png", "图 5-14 学习记忆摘要")),
                ("### 8.3 掌握度报告", fig("S15-mastery-update.png", "图 5-15 效果评估与掌握度变化")),
                ("### 11.1 教师看板", fig("S17-teacher-dashboard.png", "图 5-16 教师端只读聚合视图")),
                ("### 11.3 学生花名册", fig("S18-student-roster.png", "图 5-17 学生学情管理页面")),
            ]
            for h, b in inserts: text = insert_after(text, h, b)
        elif md == "01_项目说明书.md":
            block = "\n\n".join([
                fig("home-system.png", "图 1-1 系统首页"), fig("S04-persona-evidence.png", "图 1-2 六维学生画像与证据链"),
                fig("S05-learning-path.png", "图 1-3 个性化学习路径"), fig("S06-agent-progress.png", "图 1-4 多智能体资源生成"),
                fig("S09-oj-workbench.png", "图 1-5 OJ 工作台"), fig("S17-teacher-dashboard.png", "图 1-6 教师端只读聚合视图"),
                fig("D07-learning-loop.png", "图 1-7 学习闭环总览"),
            ])
            text = insert_after(text, "## 六、创新价值", block)
        elif md == "02_系统开发说明书.md":
            block = "\n\n".join([
                fig("D01-system-architecture.png", "图 2-1 系统总体架构"), fig("D02-agent-six-layers.png", "图 2-2 多智能体六层结构"),
                fig("D03-resource-dag.png", "图 2-3 四阶段资源 DAG"), fig("D04-oj-trace-flow.png", "图 2-4 OJ/Trace 数据流"),
                fig("D05-database-er.png", "图 2-5 数据库 E-R 图"), fig("D06-dual-safety.png", "图 2-6 内容安全与执行安全双链路"),
            ])
            text = insert_after(text, "## 一、总体架构", block)
        elif md == "06_开源与AI_Coding说明.md":
            block = "\n\n".join([
                fig("D08-dependency-categories.png", "图 6-1 第三方依赖分类"),
                fig("D09-ai-review-flow.png", "图 6-2 AI Coding 人工复核流程"),
            ]) + "\n\n```text\n仓库根目录/\n├─ LICENSE\n└─ THIRD_PARTY_LICENSES.md\n```\n\n*图 6-3 LICENSE 与 THIRD_PARTY_LICENSES 文件位置（仓库实扫）*"
            text = insert_after(text, "## 一、开源项目使用清单", block)
        p.write_text(text, encoding="utf-8")

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); tcPr.append(shd)

def set_run_font(run, east="宋体", latin="Times New Roman", size=10.5, bold=None, color=None):
    run.font.name = latin; run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east)
    run._element.rPr.rFonts.set(qn("w:ascii"), latin); run._element.rPr.rFonts.set(qn("w:hAnsi"), latin)
    run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if color: run.font.color.rgb = RGBColor.from_string(color)

def add_field(p, code):
    r=p.add_run(); fld=OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), code); r._r.addnext(fld)

def add_toc_field(p):
    r1=p.add_run()._r; begin=OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"),"begin"); r1.append(begin)
    r2=p.add_run()._r; instr=OxmlElement("w:instrText"); instr.set(qn("xml:space"),"preserve"); instr.text=' TOC \\o "1-3" \\h \\z \\u '; r2.append(instr)
    r3=p.add_run()._r; sep=OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"),"separate"); r3.append(sep)
    p.add_run("目录字段：打开文档后可右键更新")
    r4=p.add_run()._r; end=OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"),"end"); r4.append(end)

def setup(doc: Document):
    sec=doc.sections[0]; sec.page_width=Cm(21); sec.page_height=Cm(29.7)
    sec.top_margin=Cm(2.5); sec.bottom_margin=Cm(2.5); sec.left_margin=Cm(2.8); sec.right_margin=Cm(2.5)
    sec.header_distance=Cm(1.5); sec.footer_distance=Cm(1.5)
    normal=doc.styles["Normal"]; normal.font.size=Pt(10.5); normal.font.name="Times New Roman"; normal._element.rPr.rFonts.set(qn("w:eastAsia"),"宋体")
    pf=normal.paragraph_format; pf.line_spacing=1.5; pf.first_line_indent=Cm(0.74); pf.space_after=Pt(3)
    for name,size in (("Heading 1",16),("Heading 2",14),("Heading 3",12)):
        st=doc.styles[name]; st.font.name="黑体"; st._element.rPr.rFonts.set(qn("w:eastAsia"),"黑体"); st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor(31,78,71)
        st.paragraph_format.first_line_indent=Cm(0); st.paragraph_format.keep_with_next=True; st.paragraph_format.space_before=Pt(10); st.paragraph_format.space_after=Pt(6)
    cap=doc.styles["Caption"]; cap.font.name="Times New Roman"; cap._element.rPr.rFonts.set(qn("w:eastAsia"),"宋体"); cap.font.size=Pt(9); cap.font.color.rgb=RGBColor(80,80,80)
    header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_run_font(header.add_run("AlgoPilot 算法领航员｜软件杯 A3"),east="宋体",size=9,color="55706A")
    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_run_font(footer.add_run("第 "),size=9); add_field(footer,"PAGE"); set_run_font(footer.add_run(" 页 / 共 "),size=9); add_field(footer,"NUMPAGES"); set_run_font(footer.add_run(" 页"),size=9)
    doc.core_properties.author="AlgoPilot 项目团队"; doc.core_properties.last_modified_by="AlgoPilot 项目团队"; doc.core_properties.comments=""

def cover(doc,title):
    for _ in range(5): doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_run_font(p.add_run(title),east="黑体",size=24,bold=True,color="1F4E47")
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(16); set_run_font(p.add_run("AlgoPilot（算法领航员）"),east="黑体",size=18,bold=True)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_run_font(p.add_run("软件杯 A3 赛题正式提交材料"),east="宋体",size=13,color="55706A")
    for _ in range(4): doc.add_paragraph()
    rows=[("文档版本","v1.2"),("文档状态","正式提交版"),("更新日期",DATE),("对应 Commit",SHA),("编制单位","AlgoPilot 项目团队")]
    t=doc.add_table(rows=0,cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    for a,b in rows:
        c=t.add_row().cells; c[0].width=Cm(4); c[1].width=Cm(10); c[0].text=a; c[1].text=b
        for x in c: x.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_page_break()
    p=doc.add_paragraph("目录",style="Heading 1"); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    toc=doc.add_paragraph(); toc.paragraph_format.first_line_indent=Cm(0); add_toc_field(toc)
    doc.add_page_break()

def add_image(doc,path,alt):
    fp=(SUB/path).resolve() if not Path(path).is_absolute() else Path(path)
    if not fp.exists(): return
    with Image.open(fp) as im: w,h=im.size
    width=Cm(15.2); height_cm=15.2*h/w
    if height_cm>18: width=Cm(18*w/h)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.first_line_indent=Cm(0); p.paragraph_format.keep_with_next=True
    r=p.add_run(); r.add_picture(str(fp),width=width)

def add_inline(p,text):
    parts=re.split(r"(`[^`]+`|\*\*[^*]+\*\*)",text)
    for part in parts:
        if not part: continue
        if part.startswith("`"):
            r=p.add_run(part[1:-1]); set_run_font(r,east="等线",latin="Consolas",size=9); r._element.get_or_add_rPr().append(OxmlElement("w:shd")); r._element.rPr[-1].set(qn("w:fill"),"F1F3F4")
        elif part.startswith("**"):
            set_run_font(p.add_run(part[2:-2]),bold=True)
        else: set_run_font(p.add_run(part))

def markdown_to_docx(text,title,out):
    doc=Document(); setup(doc); cover(doc,title)
    lines=text.splitlines(); i=0; in_code=False; code=[]
    while i<len(lines):
        line=lines[i].rstrip()
        if line.startswith("```"):
            if not in_code: in_code=True; code=[]
            else:
                p=doc.add_paragraph(); p.paragraph_format.first_line_indent=Cm(0); p.paragraph_format.line_spacing=1.0
                pPr=p._p.get_or_add_pPr(); shd=OxmlElement("w:shd"); shd.set(qn("w:fill"),"F3F4F6"); pPr.append(shd)
                r=p.add_run("\n".join(code)); set_run_font(r,east="等线",latin="Consolas",size=9)
                in_code=False
            i+=1; continue
        if in_code: code.append(line); i+=1; continue
        if not line or line=="---": i+=1; continue
        m=re.match(r"^(#{1,3})\s+(.*)",line)
        if m:
            level=len(m.group(1)); txt=m.group(2)
            if level==1 and txt.startswith("AlgoPilot"): i+=1; continue
            doc.add_paragraph(txt,style=f"Heading {level}"); i+=1; continue
        m=re.match(r"!\[(.*?)\]\((.*?)\)",line)
        if m: add_image(doc,m.group(2),m.group(1)); i+=1; continue
        if line.startswith("*") and line.endswith("*") and line[1:-1].startswith("图 "):
            p=doc.add_paragraph(line[1:-1],style="Caption"); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.first_line_indent=Cm(0); i+=1; continue
        if line.startswith("|") and i+1<len(lines) and re.match(r"^\|?[\s:|-]+\|",lines[i+1]):
            rows=[]; i+=2
            header=[x.strip() for x in line.strip("|").split("|")]
            while i<len(lines) and lines[i].startswith("|"):
                rows.append([x.strip() for x in lines[i].strip("|").split("|")]); i+=1
            t=doc.add_table(rows=1,cols=len(header)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=True
            for j,x in enumerate(header): t.rows[0].cells[j].text=re.sub(r"\*\*|`", "", x); set_cell_shading(t.rows[0].cells[j],"DCEBE7")
            trPr=t.rows[0]._tr.get_or_add_trPr(); rep=OxmlElement("w:tblHeader"); rep.set(qn("w:val"),"true"); trPr.append(rep)
            for row in rows:
                c=t.add_row().cells
                for j in range(len(header)): c[j].text=re.sub(r"\*\*|`", "", row[j] if j<len(row) else "")
            for row in t.rows:
                for cell in row.cells:
                    cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    for p in cell.paragraphs:
                        p.paragraph_format.first_line_indent=Cm(0); p.paragraph_format.line_spacing=1.1
                        for r in p.runs: set_run_font(r,size=9,bold=(row is t.rows[0]))
            continue
        if re.match(r"^\s*[-*+]\s+",line):
            p=doc.add_paragraph(style="List Bullet"); p.paragraph_format.first_line_indent=Cm(0); add_inline(p,re.sub(r"^\s*[-*+]\s+","",line)); i+=1; continue
        if re.match(r"^\s*\d+\.\s+",line):
            p=doc.add_paragraph(style="List Number"); p.paragraph_format.first_line_indent=Cm(0); add_inline(p,re.sub(r"^\s*\d+\.\s+","",line)); i+=1; continue
        if line.startswith(">"):
            p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(0.7); p.paragraph_format.first_line_indent=Cm(0); pPr=p._p.get_or_add_pPr(); shd=OxmlElement("w:shd"); shd.set(qn("w:fill"),"EEF5F3"); pPr.append(shd); add_inline(p,line.lstrip("> ")); i+=1; continue
        p=doc.add_paragraph(); add_inline(p,line); i+=1
    if "系统开发" in title:
        doc.add_page_break()
    doc.add_paragraph("版本记录",style="Heading 1")
    t=doc.add_table(rows=2,cols=4); t.style="Table Grid"; vals=[["版本","日期","对应 Commit","说明"],["v1.2",DATE,SHA,"正式提交版：代码核验、真实截图、测试实跑与版式检查"]]
    for a,row in enumerate(vals):
        for b,v in enumerate(row): t.cell(a,b).text=v
    for row in t.rows:
        trPr=row._tr.get_or_add_trPr(); cant=OxmlElement("w:cantSplit"); cant.set(qn("w:val"),"true"); trPr.append(cant)
    doc.save(out)

def scrub(path):
    with zipfile.ZipFile(path,"r") as zin:
        items={n:zin.read(n) for n in zin.namelist() if "comments" not in n.lower() and n!="docProps/custom.xml"}
    core=items.get("docProps/core.xml",b"")
    if core:
        s=core.decode("utf-8"); s=re.sub(r"<dc:creator>.*?</dc:creator>","<dc:creator>AlgoPilot 项目团队</dc:creator>",s); s=re.sub(r"<cp:lastModifiedBy>.*?</cp:lastModifiedBy>","<cp:lastModifiedBy>AlgoPilot 项目团队</cp:lastModifiedBy>",s); items["docProps/core.xml"]=s.encode("utf-8")
    tmp=path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(tmp,"w",zipfile.ZIP_DEFLATED) as zout:
        for n,data in items.items(): zout.writestr(n,data)
    tmp.replace(path)

if __name__ == "__main__":
    update_sources()
    for md,name,title in DOCS:
        out=OUT/name; markdown_to_docx((SUB/md).read_text(encoding="utf-8"),title,out); scrub(out)
        print(out)
